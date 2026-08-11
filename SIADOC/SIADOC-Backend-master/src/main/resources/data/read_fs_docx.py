from docx import Document

file_path = r"C:\Users\HP\Documents\01 Mars\SIADOC\SIADOC\src\main\resources\data\organigramme FS.docx"

try:
    doc = Document(file_path)
    print("--- FULL TEXT ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(" | ".join(row_text))
            
except Exception as e:
    print(f"Error: {e}")
