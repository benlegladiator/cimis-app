import docx
import json

def read_docx(file_path):
    doc = docx.Document(file_path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())
    return list(dict.fromkeys(full_text))

file_path = r"C:\Users\HP\Documents\01 Mars\SIADOC\SIADOC\src\main\resources\data\shema organigramme.docx"
try:
    text = read_docx(file_path)
    print(json.dumps(text, indent=2, ensure_ascii=False))
except Exception as e:
    print(f"Error: {e}")
